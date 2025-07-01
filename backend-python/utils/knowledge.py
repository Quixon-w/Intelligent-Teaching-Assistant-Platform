from langchain_community.document_loaders import UnstructuredMarkdownLoader, TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from langchain.schema import Document
from FlagEmbedding import BGEM3FlagModel, FlagReranker
import chromadb
import os
import fitz  # PyMuPDF 用于处理 PDF 文件
import docx  # python-docx 用于处理 DOCX 文件
import re
import uuid
from typing import List, Dict, Any, Optional
from config.settings import get_settings

# 智能文本分块器
class SmartTextSplitter:
    def __init__(self, chunk_size=768, chunk_overlap=256):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def split_text(self, text, file_type="txt"):
        """根据文件类型使用不同的分块策略"""
        if file_type == "md":
            return self._split_markdown(text)
        elif file_type == "pdf":
            return self._split_pdf(text)
        else:
            return self._split_general_text(text)
    
    def _split_markdown(self, text):
        """Markdown文件的智能分块"""
        # 使用MarkdownHeaderTextSplitter按标题分块
        headers_to_split_on = [
            ("#", "标题1"),
            ("##", "标题2"),
            ("###", "标题3"),
            ("####", "标题4"),
        ]
        markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            return_each_line=False,
        )
        
        # 先按标题分块
        header_splits = markdown_splitter.split_text(text)
        
        # 如果分块太大，再按段落分块，使用重叠
        final_splits = []
        for split in header_splits:
            if len(split.page_content) > self.chunk_size * 2:
                # 按段落进一步分割，使用重叠
                paragraphs = self._split_by_paragraphs_with_overlap(split.page_content)
                final_splits.extend(paragraphs)
            else:
                final_splits.append(split)
        
        return final_splits
    
    def _split_pdf(self, text):
        """PDF文件的智能分块"""
        # 按段落和句子分块，使用重叠
        paragraphs = self._split_by_paragraphs_with_overlap(text)
        final_splits = []
        
        for para in paragraphs:
            if len(para.page_content) > self.chunk_size:
                # 按句子进一步分割，使用重叠
                sentences = self._split_by_sentences_with_overlap(para.page_content)
                final_splits.extend(sentences)
            else:
                final_splits.append(para)
        
        return final_splits
    
    def _split_general_text(self, text):
        """通用文本的智能分块"""
        # 先按段落分块，使用重叠
        paragraphs = self._split_by_paragraphs_with_overlap(text)
        final_splits = []
        
        for para in paragraphs:
            if len(para.page_content) > self.chunk_size:
                # 按句子进一步分割，使用重叠
                sentences = self._split_by_sentences_with_overlap(para.page_content)
                final_splits.extend(sentences)
            else:
                final_splits.append(para)
        
        return final_splits
    
    def _split_by_paragraphs_with_overlap(self, text):
        """按段落分割文本，使用重叠"""
        paragraphs = re.split(r'\n\s*\n', text.strip())
        docs = []
        
        if not paragraphs:
            return [Document(page_content=text)]
        
        current_chunk = ""
        overlap_text = ""
        
        for i, para in enumerate(paragraphs):
            # 如果当前段落加上重叠文本超过限制，保存当前块
            if len(current_chunk) + len(para) > self.chunk_size and current_chunk:
                docs.append(Document(page_content=current_chunk.strip()))
                
                # 保存重叠部分（最后几个段落）
                overlap_paragraphs = current_chunk.split('\n\n')[-2:]  # 保留最后2个段落作为重叠
                overlap_text = '\n\n'.join(overlap_paragraphs)
                current_chunk = overlap_text + '\n\n' + para
            else:
                current_chunk += para + "\n\n"
        
        # 添加最后一个块
        if current_chunk:
            docs.append(Document(page_content=current_chunk.strip()))
        
        return docs
    
    def _split_by_sentences_with_overlap(self, text):
        """按句子分割文本，使用重叠"""
        # 中文句子分割，保持问答完整性
        sentences = re.split(r'([。！？；])', text)
        docs = []
        
        if not sentences:
            return [Document(page_content=text)]
        
        current_chunk = ""
        overlap_sentences = []
        
        for i in range(0, len(sentences), 2):
            if i + 1 < len(sentences):
                sentence = sentences[i] + sentences[i + 1]
            else:
                sentence = sentences[i]
            
            # 如果当前句子加上重叠句子超过限制，保存当前块
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                docs.append(Document(page_content=current_chunk.strip()))
                
                # 保存重叠部分（最后几个句子）
                overlap_sentences = current_chunk.split('。')[-3:]  # 保留最后3个句子作为重叠
                overlap_text = '。'.join(overlap_sentences) + '。'
                current_chunk = overlap_text + sentence
            else:
                current_chunk += sentence
        
        # 添加最后一个块
        if current_chunk:
            docs.append(Document(page_content=current_chunk.strip()))
        
        return docs

# ChromaDB管理器
class ChromaDBManager:
    def __init__(self, host: str = "localhost", port: int = 8000):
        """初始化ChromaDB管理器"""
        self.host = host
        self.port = port
        self.client = None
        self._init_client()
    
    def _init_client(self):
        """初始化ChromaDB客户端"""
        try:
            self.client = chromadb.HttpClient(host=self.host, port=self.port)
            print(f"✅ 成功连接到ChromaDB: {self.host}:{self.port}")
        except Exception as e:
            print(f"❌ 连接ChromaDB失败: {e}")
            raise
    
    def get_or_create_collection(self, collection_name: str):
        """获取或创建collection"""
        try:
            collection = self.client.get_or_create_collection(
                name=collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            return collection
        except Exception as e:
            print(f"❌ 创建collection失败: {e}")
            raise
    
    def add_documents(self, collection_name: str, documents: List[Document], embeddings: List[List[float]]):
        """添加文档到collection"""
        try:
            collection = self.get_or_create_collection(collection_name)
            
            # 生成唯一ID
            ids = [str(uuid.uuid4()) for _ in documents]
            texts = [doc.page_content for doc in documents]
            
            # 添加文档
            collection.add(
                ids=ids,
                embeddings=embeddings,
                documents=texts
            )
            
            print(f"✅ 成功添加 {len(documents)} 个文档到collection: {collection_name}")
            return ids
        except Exception as e:
            print(f"❌ 添加文档失败: {e}")
            raise
    
    def search_documents(self, collection_name: str, query_embedding: List[float], top_k: int = 5):
        """搜索文档"""
        try:
            collection = self.client.get_collection(collection_name)
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=top_k,
                include=['documents', 'distances']
            )
            return results
        except Exception as e:
            print(f"❌ 搜索文档失败: {e}")
            raise
    
    def delete_collection(self, collection_name: str):
        """删除collection"""
        try:
            self.client.delete_collection(collection_name)
            print(f"✅ 成功删除collection: {collection_name}")
        except Exception as e:
            print(f"❌ 删除collection失败: {e}")
            raise
    
    def list_collections(self):
        """列出所有collections"""
        try:
            collections = self.client.list_collections()
            return collections
        except Exception as e:
            print(f"❌ 列出collections失败: {e}")
            raise

# BGEM3模型管理器
class BGEM3Manager:
    def __init__(self, model_path: str):
        """初始化BGEM3模型管理器"""
        self.model_path = model_path
        self.model = None
        self._load_model()
    
    def _load_model(self):
        """加载BGEM3模型"""
        try:
            self.model = BGEM3FlagModel(self.model_path, use_fp16=True)
            print(f"✅ 成功加载BGEM3模型: {self.model_path}")
        except Exception as e:
            print(f"❌ 加载BGEM3模型失败: {e}")
            raise
    
    def encode(self, texts: List[str]) -> List[List[float]]:
        """编码文本为向量"""
        try:
            embeddings = self.model.encode(
                texts,
                max_length=512
            )["dense_vecs"].tolist()
            return embeddings
        except Exception as e:
            print(f"❌ 编码文本失败: {e}")
            raise

# BGE-Reranker模型管理器
class BGERerankerManager:
    def __init__(self, model_path: str):
        """初始化BGE-Reranker模型管理器"""
        self.model_path = model_path
        self.model = None
        self._load_model()
    
    def _load_model(self):
        """加载BGE-Reranker模型"""
        try:
            self.model = FlagReranker(self.model_path, use_fp16=True)
            print(f"✅ 成功加载BGE-Reranker模型: {self.model_path}")
        except Exception as e:
            print(f"❌ 加载BGE-Reranker模型失败: {e}")
            raise
    
    def compute_score(self, text_a: str, text_b: str) -> float:
        """计算两个文本的相关性分数"""
        try:
            score = self.model.compute_score([text_a, text_b])
            return score
        except Exception as e:
            print(f"❌ 计算相关性分数失败: {e}")
            raise
    
    def rerank_documents(self, query: str, documents: List[str]) -> List[tuple]:
        """重排序文档"""
        try:
            scores = []
            for doc in documents:
                score = self.compute_score(query, doc)
                scores.append((score, doc))
            
            # 按分数降序排序
            scores.sort(key=lambda x: x[0], reverse=True)
            return scores
        except Exception as e:
            print(f"❌ 重排序文档失败: {e}")
            raise

# 加载 DOCX 文件
def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)

# 加载 PDF 文件
def load_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# 加载文件
def load_documents(dir_path):
    docs = []
    print(f"正在加载目录: {dir_path}")  # 调试输出
    
    # 使用配置管理系统获取分块参数
    settings = get_settings()
    smart_splitter = SmartTextSplitter(
        chunk_size=settings.VECTOR_DB_CHUNK_SIZE, 
        chunk_overlap=settings.VECTOR_DB_CHUNK_OVERLAP
    )
    
    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)
        file_type = filename.split('.')[-1].lower()
        
        try:
            if file_type == 'md':
                print(f"加载 Markdown 文件: {filename}")  # 调试输出
                loader = UnstructuredMarkdownLoader(file_path)
                raw_docs = loader.load()
                # 使用智能分块
                for doc in raw_docs:
                    split_docs = smart_splitter.split_text(doc.page_content, "md")
                    docs.extend(split_docs)
                    
            elif file_type == 'txt':
                print(f"加载文本文件: {filename}")  # 调试输出
                loader = TextLoader(file_path)
                raw_docs = loader.load()
                # 使用智能分块
                for doc in raw_docs:
                    split_docs = smart_splitter.split_text(doc.page_content, "txt")
                    docs.extend(split_docs)
                    
            elif file_type == 'pdf':
                print(f"加载 PDF 文件: {filename}")  # 调试输出
                pdf_text = load_pdf(file_path)
                # 使用智能分块
                split_docs = smart_splitter.split_text(pdf_text, "pdf")
                docs.extend(split_docs)
                
            elif file_type == 'docx':
                print(f"加载 DOCX 文件: {filename}")  # 调试输出
                docx_text = load_docx(file_path)
                # 使用智能分块
                split_docs = smart_splitter.split_text(docx_text, "txt")
                docs.extend(split_docs)
                
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")
            continue
    
    print(f"总共加载了 {len(docs)} 个文档块")  # 调试输出
    return docs

# 创建向量知识库
def create_vector_db(documents, collection_name):
    if not documents:
        print("没有文档需要处理")
        return None

    try:
        # 使用配置管理系统获取模型路径和ChromaDB配置
        settings = get_settings()
        
        # 初始化BGEM3模型
        bgem3_manager = BGEM3Manager(str(settings.BGEM3_MODEL_PATH))
        
        # 初始化ChromaDB管理器
        chroma_manager = ChromaDBManager(
            host=settings.CHROMADB_HOST,
            port=settings.CHROMADB_PORT
        )
        
        # 生成文档向量
        texts = [doc.page_content for doc in documents]
        embeddings = bgem3_manager.encode(texts)
        
        # 添加到ChromaDB
        ids = chroma_manager.add_documents(collection_name, documents, embeddings)
        
        print(f"✅ 成功创建向量知识库，collection: {collection_name}")
        print(f"✅ 添加了 {len(documents)} 个文档块")
        return chroma_manager
        
    except Exception as e:
        print(f"❌ 创建向量知识库失败: {e}")
        return None

def update_knowledge_db(userId, isTeacher=False, courseID=None, lessonNum=None, isResource=False, isAsk=False):
    """
    更新知识库
    :param userId: 用户ID
    :param isTeacher: 是否为教师模式
    :param courseID: 课程ID（教师模式下非ask文件时必填）
    :param lessonNum: 课时号（教师模式下非ask文件时必填）
    :param isResource: 是否为学习资料
    :param isAsk: 是否为可提问文件
    """
    try:
        # 使用配置管理系统获取路径
        settings = get_settings()
        
        # 根据isTeacher和文件类型决定存储路径，与upload.py保持一致
        if isTeacher:
            # 教师模式
            base_folder = str(settings.TEACHERS_DIR)
            if isResource:
                # 学习资料：保存到courseId级别
                session_folder = f"{base_folder}/{userId}/{courseID}"
            elif isAsk:
                # 可对文件进行提问的文件：保存在ask文件夹
                session_folder = f"{base_folder}/{userId}/ask"
            else:
                # 大纲与习题生成参考文件：保存到lessonNum级别
                session_folder = f"{base_folder}/{userId}/{courseID}/{lessonNum}"
        else:
            # 学生模式：存储在Students目录下的userId文件夹中
            base_folder = str(settings.STUDENTS_DIR)
            if isAsk:
                # 学生上传的可提问文件：保存在ask文件夹
                session_folder = f"{base_folder}/{userId}/ask"
            else:
                # 其他文件：保存在userId文件夹
                session_folder = f"{base_folder}/{userId}"
        
        print(f"开始更新知识库，userId: {userId}, isTeacher: {isTeacher}, courseID: {courseID}, lessonNum: {lessonNum}, isResource: {isResource}, isAsk: {isAsk}")  # 调试输出

        # 加载文档
        docs = load_documents(session_folder)
        if not docs:
            print("没有找到任何支持的文档类型，请确认上传的文件类型。")  # 调试输出
            return None
        
        # 生成collection名称
        collection_name = f"kb_{userId}_{courseID or 'student'}_{lessonNum or 'default'}"
        if isAsk:
            collection_name += "_ask"
        
        # 创建向量知识库
        chroma_manager = create_vector_db(docs, collection_name)
        return chroma_manager
        
    except ValueError as e:
        print(f"路径参数错误: {e}")
        return None
    except Exception as e:
        print(f"更新知识库时出错: {e}")
        return None

# 加载已存在的向量数据库
def load_vector_db(userId, isTeacher=False, courseID=None, lessonNum=None):
    """
    加载已存在的向量数据库
    :param userId: 用户ID
    :param isTeacher: 是否为教师模式
    :param courseID: 课程ID（教师模式下必填）
    :param lessonNum: 课时号（教师模式下必填）
    """
    try:
        # 使用配置管理系统获取ChromaDB配置
        settings = get_settings()
        
        # 生成collection名称
        collection_name = f"kb_{userId}_{courseID or 'student'}_{lessonNum or 'default'}"
        
        # 初始化ChromaDB管理器
        chroma_manager = ChromaDBManager(
            host=settings.CHROMADB_HOST,
            port=settings.CHROMADB_PORT
        )
        
        # 检查collection是否存在
        try:
            collection = chroma_manager.client.get_collection(collection_name)
            print(f"✅ 成功加载向量数据库，collection: {collection_name}")
            return chroma_manager
        except Exception as e:
            print(f"❌ 向量数据库不存在: {collection_name}")
            return None
                
    except Exception as e:
        print(f"加载向量数据库时出错: {e}")
        return None

# 搜索知识库
def search_knowledge_db(query: str, chroma_manager: ChromaDBManager, collection_name: str, 
                       top_k: int = 5, use_rerank: bool = True):
    """
    搜索知识库
    :param query: 查询文本
    :param chroma_manager: ChromaDB管理器
    :param collection_name: collection名称
    :param top_k: 返回结果数量
    :param use_rerank: 是否使用重排序
    """
    try:
        settings = get_settings()
        
        # 初始化BGEM3模型
        bgem3_manager = BGEM3Manager(str(settings.BGEM3_MODEL_PATH))
        
        # 生成查询向量
        query_embedding = bgem3_manager.encode([query])[0]
        
        # 搜索文档
        results = chroma_manager.search_documents(collection_name, query_embedding, top_k * 2)  # 获取更多候选
        
        if not results['documents'] or not results['documents'][0]:
            print("没有找到相关文档")
            return []
        
        documents = results['documents'][0]
        distances = results['distances'][0]
        
        # 如果启用重排序
        if use_rerank and len(documents) > 1:
            try:
                # 初始化BGE-Reranker模型
                reranker_manager = BGERerankerManager(str(settings.BGE_RERANKER_MODEL_PATH))
                
                # 重排序
                reranked_results = reranker_manager.rerank_documents(query, documents)
                
                # 返回前top_k个结果
                final_results = []
                for score, doc in reranked_results[:top_k]:
                    final_results.append({
                        'document': doc,
                        'score': score,
                        'reranked': True
                    })
                
                print(f"✅ 重排序完成，返回 {len(final_results)} 个结果")
                return final_results
                
            except Exception as e:
                print(f"重排序失败，使用原始结果: {e}")
                # 如果重排序失败，使用原始结果
                pass
        
        # 不使用重排序或重排序失败时，使用原始结果
        final_results = []
        for i, (doc, distance) in enumerate(zip(documents, distances)):
            if i >= top_k:
                break
            final_results.append({
                'document': doc,
                'score': 1.0 - distance,  # 距离转换为相似度分数
                'reranked': False
            })
        
        print(f"✅ 搜索完成，返回 {len(final_results)} 个结果")
        return final_results
        
    except Exception as e:
        print(f"搜索知识库失败: {e}")
        return []