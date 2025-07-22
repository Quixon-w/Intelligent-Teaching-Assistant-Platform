# -*- coding: utf-8 -*-
import os
import json
from fastapi import APIRouter, HTTPException, status, Request
from pydantic import BaseModel, Field
from typing import Union, Optional, List, Dict, Any
import asyncio
from threading import Lock
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils.rwkv import *
from utils.knowledge import load_vector_db, search_knowledge_db, ChromaDBManager, search_domain_knowledge, build_enhanced_prompt
import global_var
from config.settings import get_settings

router = APIRouter()

# 全局锁，用于控制并发请求
create_lock = Lock()


class CreateContentDesignBody(BaseModel):
    user_id: Union[str, int] = Field(..., description="用户ID，用于确定存储路径")
    session_id: str = Field(..., description="会话ID")
    course_id: str = Field(..., description="课程ID")
    lesson_num: str = Field(..., description="课时号，必填")
    is_teacher: bool = Field(False, description="是否为教师用户")
    # 教学设计字数控制：课时教学设计控制在800-1200字
    max_words: int = Field(1000, description="最大字数限制，课时教学设计建议1000字", ge=300, le=2000)

    model_config = {
        "json_schema_extra": {
            "example": {
                "user_id": "teacher123",
                "session_id": "session456",
                "course_id": "math101",
                "lesson_num": "lesson01",
                "is_teacher": True,
                "max_words": 1000
            }
        }
    }
    
    def __init__(self, **data):
        super().__init__(**data)
        # 确保user_id是字符串类型
        if isinstance(self.user_id, int):
            self.user_id = str(self.user_id)


def get_user_path(user_id: str, is_teacher: bool) -> str:
    """根据userID和isTeacher确定用户路径"""
    settings = get_settings()
    if is_teacher:
        base_dir = settings.TEACHERS_DIR
    else:
        base_dir = settings.STUDENTS_DIR
    return os.path.join(str(base_dir), user_id)


def get_lesson_content_from_chromadb(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> str:
    """
    从ChromaDB知识库获取课时内容，返回合并的文本内容
    """
    print(f"正在从ChromaDB获取课时内容: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, isTeacher={is_teacher}")
    
    try:
        # 生成collection名称，与update_knowledge_db保持一致
        collection_name = f"kb_{user_id}_{course_id}_{lesson_num}"
        
        # 初始化ChromaDB管理器
        from utils.knowledge import ChromaDBManager
        from config.settings import get_settings
        
        settings = get_settings()
        chroma_manager = ChromaDBManager(
            host=settings.CHROMADB_HOST,
            port=settings.CHROMADB_PORT
        )
        
        # 获取collection中的所有文档
        try:
            collection = chroma_manager.client.get_collection(collection_name)
            
            # 获取所有文档
            results = collection.get(include=['documents'])
            
            if results and results['documents']:
                text_blocks = results['documents']
                # 合并文本块
                combined_content = "\n\n".join(text_blocks)
                print(f"成功从ChromaDB获取课时内容，总长度: {len(combined_content)} 字符")
                return combined_content
            else:
                print("ChromaDB中没有找到文档")
                return None
        
        except Exception as e:
            print(f"从ChromaDB获取文档失败: {e}")
            return None
            
    except Exception as e:
        print(f"加载ChromaDB知识库失败: {e}")
        return None


def get_file_content(user_id: str, course_id: str, lesson_num: str, is_teacher: bool):
    """
    获取课时内容，优先从ChromaDB知识库获取，如果失败则从文件获取
    """
    print(f"正在获取课时内容: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, isTeacher={is_teacher}")
    
    # 首先尝试从ChromaDB知识库获取内容
    content = get_lesson_content_from_chromadb(user_id, course_id, lesson_num, is_teacher)
    if content:
        print("成功从ChromaDB知识库获取课时内容")
        return content
    
    # 如果ChromaDB获取失败，回退到文件读取方式
    print("ChromaDB知识库获取失败，回退到文件读取方式")
    
    # 获取用户路径
    user_path = get_user_path(user_id, is_teacher)
    
    # 搜索特定课时的内容
    lesson_folder = os.path.join(user_path, course_id, lesson_num)
    
    # 读取特定课时的文件内容
    if not os.path.exists(lesson_folder):
        print(f"课时目录不存在: {lesson_folder}")
        return None
    
    content = read_files_in_folder(lesson_folder)
    return content


def read_files_in_folder(folder_path: str) -> str:
    """
    读取文件夹中的所有支持的文件内容
    """
    if not os.path.exists(folder_path):
        print(f"文件夹不存在: {folder_path}")
        return None
    
    content_parts = []
    
    try:
        # 遍历文件夹中的所有文件
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            # 跳过目录和隐藏文件
            if os.path.isdir(file_path) or filename.startswith('.'):
                continue
            
            # 只处理支持的文件类型
            file_extension = os.path.splitext(filename)[1].lower()
            if file_extension not in ['.txt', '.md', '.pdf', '.docx']:
                continue
            
            try:
                if file_extension in ['.txt', '.md']:
                    # 读取文本文件
                    with open(file_path, 'r', encoding='utf-8') as f:
                        file_content = f.read()
                        content_parts.append(f"文件: {filename}\n{file_content}")
                elif file_extension in ['.pdf', '.docx']:
                    # 对于PDF和DOCX文件，这里需要添加相应的解析逻辑
                    # 暂时跳过，或者可以调用外部工具进行解析
                    content_parts.append(f"文件: {filename} (PDF/DOCX文件，需要特殊处理)")
                    
            except Exception as e:
                print(f"读取文件 {filename} 时出错: {e}")
                continue
        
        if content_parts:
            return "\n\n".join(content_parts)
        else:
            print(f"文件夹 {folder_path} 中没有找到可读取的文件")
            return None
            
    except Exception as e:
        print(f"读取文件夹 {folder_path} 时出错: {e}")
        return None


def generate_content_design_prompt(content: str, max_words: int) -> str:
    """
    根据内容生成教学内容设计提示词
    """
    prompt = f"""请根据以下教学内容，设计一个详细且丰富的教学方案。要求：\n\n1. 结构清晰，层次分明，使用数字编号（如1.1、1.2、2.1等）\n2. 包含教学目标、重点难点、知识讲解、实训练习与指导、建议时间分布、教学方法等完整内容\n3. 每一部分都要有详细的描述和具体的教学建议，内容要尽量丰富、具体、可操作\n4. 字数控制在{max_words}字左右，内容尽量详细，力求达到或超过目标字数\n5. 适合实际教学使用，便于教师直接参考和应用\n6. 体现教学设计的专业性和创新性\n\n请生成完整的教学设计方案，确保内容详细且可操作：\n\n教学内容：\n{content[:2000]}  # 限制内容长度，避免token过多\n\n教学设计："""
    return prompt


def save_content_design_to_file(user_id: str, course_id: str, lesson_num: str, content_design: str, is_teacher: bool) -> dict:
    """
    将生成的教学设计保存为DOCX文件
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 创建教学设计保存目录
        content_design_dir = os.path.join(user_path, course_id, lesson_num, "content_design")
        os.makedirs(content_design_dir, exist_ok=True)
        
                # 生成文件名（包含时间戳）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"content_design_{timestamp}.docx"
        file_path = os.path.join(content_design_dir, filename)
        
        # 创建DOCX文档
        doc = Document()
        
        # 添加标题
        title = f"教学设计 - {course_id} - {lesson_num}"
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Inches(0.2)  # 设置字体大小
        title_run.font.bold = True  # 设置粗体
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
        
        # 添加空行
        doc.add_paragraph()
        
        # 添加教学设计内容
        lines = content_design.split('\n')
        for line in lines:
            line = line.strip()
            if line:
                # 处理标题行（以数字开头的行）
                if line[0].isdigit() and '.' in line[:3]:
                    # 标题行
                    heading_paragraph = doc.add_paragraph()
                    heading_run = heading_paragraph.add_run(line)
                    heading_run.font.size = Inches(0.14)  # 标题字体大小
                    heading_run.font.bold = True  # 标题粗体
                else:
                    # 普通内容行
                    content_paragraph = doc.add_paragraph()
                    content_run = content_paragraph.add_run(line)
                    content_run.font.size = Inches(0.12)  # 正文字体大小
            else:
                # 空行
                doc.add_paragraph()
        
        # 保存DOCX文件
        doc.save(file_path)
        print(f"教学设计DOCX已保存到: {file_path}")
        
        return {
            "success": True,
            "file_path": file_path,
            "filename": filename,
            "download_url": f"/v1/download/content_design/{user_id}/{course_id}/{lesson_num}/{filename}"
        }
        
    except Exception as e:
        print(f"保存教学设计文件时出错: {e}")
        return {
            "success": False,
            "error": str(e)
        }


async def generate_content_design_with_rwkv(prompt: str, request: Request, max_words: int = 1000):
    """
    使用RWKV模型生成教学设计
    """
    print("开始使用RWKV模型生成教学设计...")
    
    try:
        # 获取RWKV模型实例
        model = global_var.get(global_var.Model)
        if model is None:
            raise Exception("RWKV模型未初始化")
        
        # 设置生成参数
        model.temperature = 0.8  # 稍微提高创造性
        model.top_p = 0.95  # 提高采样范围
        
        # 重要：临时增加max_tokens_per_generation以支持更长的生成
        original_max_tokens = model.max_tokens_per_generation
        max_tokens = 4000  # 降低到4000，避免卡住
        model.max_tokens_per_generation = max_tokens
        
        # 生成教学设计内容
        content_design = ""
        token_count = 0
        min_words = max(400, max_words * 0.4)  # 调整最小字数要求
        
        print("开始生成教学设计...")
        print(f"提示词长度: {len(prompt)} 字符")
        print(f"最大token数: {max_tokens}")
        print(f"目标字数: {max_words}")
        print(f"最小字数: {min_words}")
        print(f"模型原始max_tokens_per_generation: {original_max_tokens}")
        print(f"临时设置为: {model.max_tokens_per_generation}")
        
        # 设置停止条件
        stop_sequences = ["结束", "完毕"]
        
        print("开始生成循环...")
        for response, delta, _, _ in model.generate(prompt, stop=stop_sequences):
            content_design += delta
            token_count += 1
            
            # 每100个token打印一次进度
            if token_count % 100 == 0:
                current_words = len([c for c in content_design if '\u4e00' <= c <= '\u9fff'])
                print(f"已生成 {token_count} tokens, 当前内容长度: {len(content_design)} 字符, 字数: {current_words}")
            
            # 检查请求是否断开
            if await request.is_disconnected():
                print("请求已断开")
                break
            
            # 检查是否达到最大token数
            if token_count >= max_tokens:
                print(f"达到最大token数: {max_tokens}")
                break
            
            # 检查字数限制 - 移除这个限制，让AI尽可能多生成
            # current_words = len([c for c in content_design if '\u4e00' <= c <= '\u9fff'])
            # if current_words >= max_words * 1.2:  # 允许超出20%的字数
            #     print(f"达到字数限制 {current_words} >= {max_words * 1.2}")
            #     break
            
        # 恢复原始设置
        model.max_tokens_per_generation = original_max_tokens
        
        # 确保句子完整性
        content_design = ensure_sentence_completeness(content_design)
        
        # 检查生成的内容是否达到最小字数要求
        final_word_count = len([c for c in content_design if '\u4e00' <= c <= '\u9fff'])
        
        if final_word_count < min_words:
            print(f"生成的内容字数不足（{final_word_count} < {min_words}），尝试重新生成...")
            # 如果字数不足，尝试重新生成一次 - 简化重试提示词，避免过长
            retry_prompt = f"请继续生成更多内容，当前只生成了{final_word_count}字，需要至少{min_words}字。请详细展开每个知识点，包含案例和练习。"
            return await generate_content_design_with_rwkv_retry(retry_prompt, request, max_words, model, max_tokens)
        
        print(f"教学设计生成完成，生成长度: {len(content_design)} 字符，字数: {final_word_count}")
        print(f"内容预览: {content_design[:200]}...")
        
        return content_design.strip()
        
    except Exception as e:
        print(f"RWKV生成教学设计时出错: {e}")
        # 确保在出错时也恢复原始设置
        try:
            model = global_var.get(global_var.Model)
            if model:
                model.max_tokens_per_generation = original_max_tokens
        except:
            pass
        raise e


async def generate_content_design_with_rwkv_retry(prompt: str, request: Request, max_words: int, model, max_tokens):
    """
    重试生成教学设计（当第一次生成字数不足时）
    """
    try:
        print("开始重试生成教学设计...")
        
        # 临时设置max_tokens_per_generation
        original_max_tokens = model.max_tokens_per_generation
        model.max_tokens_per_generation = 3000  # 重试时使用更小的token数
        
        content_design = ""
        token_count = 0
        
        # 设置停止条件 - 减少停止条件，让AI继续生成
        stop_sequences = ["完毕"]
        
        for response, delta, _, _ in model.generate(prompt, stop=stop_sequences):
            content_design += delta
            token_count += 1
            
            # 计算当前字数（中文字符）
            current_words = len([c for c in content_design if '\u4e00' <= c <= '\u9fff'])
            
            # 检查token数量限制
            if token_count >= max_tokens:
                print(f"重试达到最大token限制 {max_tokens}")
                break
            
            # 检查字数限制 - 移除这个限制，让AI尽可能多生成
            # if current_words >= max_words * 1.5:  # 重试时允许更多字数
            #     print(f"重试达到字数限制 {current_words} >= {max_words * 1.5}")
            #     break
            
            # 检查请求是否断开
            if await request.is_disconnected():
                print("请求已断开")
                break
        
        # 恢复原始设置
        model.max_tokens_per_generation = original_max_tokens
        
        # 确保句子完整性
        content_design = ensure_sentence_completeness(content_design)
        
        final_word_count = len([c for c in content_design if '\u4e00' <= c <= '\u9fff'])
        print(f"重试教学设计生成完成，生成长度: {len(content_design)} 字符，字数: {final_word_count}")
        
        return content_design.strip()
        
    except Exception as e:
        print(f"重试生成教学设计时出错: {e}")
        # 确保在出错时也恢复原始设置
        try:
            model.max_tokens_per_generation = original_max_tokens
        except:
            pass
        raise e


def ensure_sentence_completeness(text: str) -> str:
    """
    确保句子完整性，避免句子被截断
    """
    if not text:
        return text
    
    # 常见的句子结束标点
    sentence_endings = ['。', '！', '？', '；', '：', '\n']
    
    # 如果文本以句子结束标点结尾，直接返回
    if text[-1] in sentence_endings:
        return text
    
    # 查找最后一个完整的句子
    last_complete_sentence = ""
    for i in range(len(text) - 1, -1, -1):
        if text[i] in sentence_endings:
            last_complete_sentence = text[:i + 1]
            break
    
    # 如果找到了完整句子，返回完整句子部分
    if last_complete_sentence:
        return last_complete_sentence
    
    # 如果没有找到完整句子，尝试在最后一个逗号处截断
    last_comma = text.rfind('，')
    if last_comma > 0:
        return text[:last_comma + 1]
    
    # 如果都没有，返回原文本
    return text


@router.post("/v1/create/content_design", tags=["Create"])
async def create_content_design(body: CreateContentDesignBody, request: Request):
    """
    创建教学设计
    """
    # 检查模型是否加载
    model: TextRWKV = global_var.get(global_var.Model)
    if model is None:
        raise HTTPException(
            status_code=status.HTTP_503_SERVICE_UNAVAILABLE,
            detail="模型未加载"
        )
    
    # 检查是否已有生成任务在进行
    if create_lock.locked():
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="已有教学设计生成任务在进行中，请稍后再试"
        )
    
    try:
        with create_lock:
            print(f"开始为课时 {body.lesson_num} 生成教学设计")
            
            # 获取课时内容
            content = get_file_content(body.user_id, body.course_id, body.lesson_num, body.is_teacher)
            if not content:
                raise HTTPException(
                    status_code=status.HTTP_404_NOT_FOUND,
                    detail=f"未找到课时 {body.lesson_num} 的内容，请先上传相关文件"
                )
            
            # 从domain知识库搜索相关内容
            domain_knowledge = search_domain_knowledge(
                query=f"教学设计 {content[:200]}",  # 使用内容前200字符作为查询
                top_k=3
            )
            
            # 构建增强的教学设计生成提示词
            enhanced_prompt = build_enhanced_prompt(
                query=f"生成课时 {body.lesson_num} 的教学设计",
                user_content=content,
                domain_knowledge=domain_knowledge,
                task_type="outline"
            )
            
            # 添加字数限制到提示词
            enhanced_prompt += f"\n\n要求：字数控制在{body.max_words}字左右，包含知识讲解、实训练习、时间分配等详细内容。"
            
            # 强制要求AI生成更多内容
            enhanced_prompt += f"\n\n重要提醒：请务必生成详细完整的内容，不要提前结束，目标字数{body.max_words}字，实际生成内容应该接近或超过这个字数。每个知识点都要详细展开，包含具体案例和练习。"
            
            # 使用RWKV模型生成教学设计
            content_design = await generate_content_design_with_rwkv(enhanced_prompt, request, body.max_words)
            
            # 保存教学设计到文件
            save_result = save_content_design_to_file(body.user_id, body.course_id, body.lesson_num, content_design, body.is_teacher)
            
            if save_result["success"]:
                return {
                    "success": True,
                    "message": "教学设计生成成功",
                    "user_id": body.user_id,
                    "session_id": body.session_id,
                    "course_id": body.course_id,
                    "lesson_num": body.lesson_num,
                    "is_teacher": body.is_teacher,
                    "content_design": content_design[:500] + "..." if len(content_design) > 500 else content_design,
                    "domain_knowledge_used": bool(domain_knowledge),
                    "download_url": save_result["download_url"],
                    "filename": save_result["filename"]
                }
            else:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail=f"保存教学设计文件失败: {save_result['error']}"
                )
                
    except HTTPException:
        raise
    except Exception as e:
        print(f"创建教学设计时出错: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"创建教学设计失败: {str(e)}"
        )


@router.get("/v1/create/content_design/status", tags=["Create"])
async def get_content_design_status(user_id: str, course_id: str, lesson_num: str, is_teacher: bool = False):
    """
    获取教学设计生成状态
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 检查教学设计目录是否存在
        content_design_dir = os.path.join(user_path, course_id, lesson_num, "content_design")
        
        if not os.path.exists(content_design_dir):
            return {
                "has_content_design": False,
                "message": "教学设计目录不存在"
            }
        
        # 检查是否有教学设计文件（支持DOCX和文本文件）
        content_design_files = [f for f in os.listdir(content_design_dir) if f.startswith('content_design_') and (f.endswith('.docx') or f.endswith('.txt'))]
        
        if content_design_files:
            # 获取最新的教学设计文件
            latest_file = max(content_design_files, key=lambda x: os.path.getctime(os.path.join(content_design_dir, x)))
            file_path = os.path.join(content_design_dir, latest_file)
            
            return {
                "success": True,
                "message": "获取教学设计状态成功",
                "user_id": user_id,
                "course_id": course_id,
                "lesson_num": lesson_num,
                "is_teacher": is_teacher,
                "has_content_design": True,
                "latest_file": latest_file,
                "file_type": "DOCX" if latest_file.endswith('.docx') else "TXT",
                "download_url": f"/v1/download/content_design/{user_id}/{course_id}/{lesson_num}/{latest_file}",
                "created_time": datetime.fromtimestamp(os.path.getctime(file_path)).strftime('%Y-%m-%d %H:%M:%S')
            }
        else:
            return {
                "has_content_design": False,
                "message": "教学设计文件不存在"
            }
            
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取教学设计状态失败: {str(e)}"
        ) 