# -*- coding: utf-8 -*-
import os
import json
from fastapi import APIRouter, HTTPException, status, Request, Response
from pydantic import BaseModel, Field
from typing import Union, Optional, List, Dict, Any
import asyncio
from threading import Lock
from datetime import datetime
import random
from docx import Document
from docx.shared import Inches

from utils.rwkv import *
from utils.knowledge import load_vector_db, search_knowledge_db, ChromaDBManager
import global_var
from config.settings import get_settings

router = APIRouter()

# 全局锁，用于控制并发请求
exercise_lock = Lock()


class ExerciseBody(BaseModel):
    user_id: Union[str, int] = Field(..., description="用户ID，用于确定存储路径")
    session_id: str = Field(..., description="会话ID")
    course_id: str = Field(..., description="课程ID")
    lesson_num: str = Field(..., description="课时号，必填")
    is_teacher: bool = Field(False, description="是否为教师用户")
    question_count: int = Field(5, description="生成题目数量", ge=1, le=20)
    difficulty: str = Field("medium", description="题目难度：easy(简单)、medium(中等)、hard(困难)")
    max_tokens: int = Field(2000, description="生成回答的最大token数", ge=500, le=4000)
    temperature: float = Field(0.7, description="生成温度", ge=0.1, le=1.0)
    generation_mode: str = Field("block", description="生成模式：block(按文本块生成)、whole(整体内容生成)")

    model_config = {
        "json_schema_extra": {
            "example": {
                "user_id": "teacher123",
                "session_id": "session456",
                "course_id": "math101",
                "lesson_num": "lesson01",
                "is_teacher": True,
                "question_count": 5,
                "difficulty": "medium",
                "max_tokens": 2000,
                "temperature": 0.7,
                "generation_mode": "block"
            }
        }
    }

    def __init__(self, **data):
        super().__init__(**data)
        # 确保user_id是字符串类型
        if isinstance(self.user_id, int):
            self.user_id = str(self.user_id)


class ExerciseQuestion(BaseModel):
    question_id: str = Field(..., description="题目ID")
    question_text: str = Field(..., description="题干")
    options: List[str] = Field(..., description="选项列表")
    correct_answer: str = Field(..., description="正确答案")
    explanation: str = Field(..., description="解析")
    knowledge_point: str = Field(..., description="所属知识点")
    difficulty: str = Field(..., description="难度等级")


class ExerciseResponse(BaseModel):
    success: bool = Field(..., description="是否成功")
    message: str = Field(..., description="响应消息")
    data: Optional[str] = Field(None, description="生成的习题原始文本")
    total_count: int = Field(0, description="生成的题目总数")
    generation_time: float = Field(0.0, description="生成耗时(秒)")


def get_user_path(user_id: str, is_teacher: bool) -> str:
    """根据userID和isTeacher确定用户路径"""
    settings = get_settings()
    if is_teacher:
        base_dir = settings.TEACHERS_DIR
    else:
        base_dir = settings.STUDENTS_DIR
    return os.path.join(str(base_dir), user_id)


def get_lesson_content_from_chromadb(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> List[str]:
    """
    从ChromaDB知识库获取课时内容，返回文本块列表
    """
    print(f"正在从ChromaDB获取课时内容: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, isTeacher={is_teacher}")
    
    try:
        # 加载向量数据库
        chroma_manager = load_vector_db(
            userId=user_id,
            isTeacher=is_teacher,
            courseID=course_id,
            lessonNum=lesson_num,
            isAsk=False      # 不是ask文件
        )
        
        if not chroma_manager:
            print("ChromaDB知识库不存在")
            return None
        
        # 生成collection名称
        collection_name = f"kb_{user_id}_{course_id}_{lesson_num}"
        
        # 获取collection中的所有文档
        try:
            collection = chroma_manager.client.get_collection(collection_name)
            
            # 获取所有文档（这里我们需要获取所有文档，所以使用一个通用的查询）
            # 由于ChromaDB没有直接获取所有文档的API，我们使用一个技巧
            # 先获取一个文档的embedding，然后用它来查询所有文档
            results = collection.get(include=['documents'])
            
            if results and results['documents']:
                text_blocks = results['documents']
                print(f"成功从ChromaDB获取 {len(text_blocks)} 个文本块")
                return text_blocks
            else:
                print("ChromaDB中没有找到文档")
                return None
                
        except Exception as e:
            print(f"从ChromaDB获取文档失败: {e}")
            return None
            
    except Exception as e:
        print(f"加载ChromaDB知识库失败: {e}")
        return None


def get_lesson_content_whole_from_chromadb(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> str:
    """
    从ChromaDB知识库获取课时内容，返回合并的文本内容
    """
    print(f"正在从ChromaDB获取课时内容: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, isTeacher={is_teacher}")
    
    text_blocks = get_lesson_content_from_chromadb(user_id, course_id, lesson_num, is_teacher)
    
    if text_blocks:
        # 合并文本块
        combined_content = "\n\n".join(text_blocks)
        print(f"成功从ChromaDB获取课时内容，总长度: {len(combined_content)} 字符")
        return combined_content
    
    print("无法从ChromaDB获取课时内容")
    return None


def get_lesson_content_fallback(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> List[str]:
    """
    备用方案：从文件系统获取课时内容
    """
    print(f"使用备用方案从文件系统获取课时内容: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, isTeacher={is_teacher}")
    
    # 获取用户路径
    user_path = get_user_path(user_id, is_teacher)
    
    # 构建文件路径
    lesson_path = os.path.join(user_path, course_id, lesson_num)
    if not os.path.exists(lesson_path):
        print(f"课时路径不存在: {lesson_path}")
        return None
    
    # 读取课时文件夹中的所有文件
    text_blocks = read_files_as_blocks(lesson_path)
    if text_blocks:
        print("成功从文件系统获取课时内容")
        return text_blocks
    
    print("无法从文件系统获取课时内容")
    return None


def get_lesson_content_whole_fallback(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> str:
    """
    备用方案：从文件系统获取课时内容，返回合并的文本
    """
    text_blocks = get_lesson_content_fallback(user_id, course_id, lesson_num, is_teacher)
    
    if text_blocks:
        # 合并文本块
        combined_content = "\n\n".join(text_blocks)
        print(f"成功从文件系统获取课时内容，总长度: {len(combined_content)} 字符")
        return combined_content
    
    return None


def get_lesson_content(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> List[str]:
    """
    获取课时内容，优先从ChromaDB获取，失败则从文件系统获取
    返回文本块列表
    """
    # 首先尝试从ChromaDB获取
    text_blocks = get_lesson_content_from_chromadb(user_id, course_id, lesson_num, is_teacher)
    if text_blocks:
        return text_blocks
    
    # 如果ChromaDB获取失败，使用备用方案
    return get_lesson_content_fallback(user_id, course_id, lesson_num, is_teacher)


def get_lesson_content_whole(user_id: str, course_id: str, lesson_num: str, is_teacher: bool) -> str:
    """
    获取课时内容，优先从ChromaDB获取，失败则从文件系统获取
    返回合并的文本内容
    """
    # 首先尝试从ChromaDB获取
    combined_content = get_lesson_content_whole_from_chromadb(user_id, course_id, lesson_num, is_teacher)
    if combined_content:
        return combined_content
    
    # 如果ChromaDB获取失败，使用备用方案
    return get_lesson_content_whole_fallback(user_id, course_id, lesson_num, is_teacher)


def read_files_as_blocks(folder_path: str) -> List[str]:
    """
    读取文件夹中的所有文本文件内容，返回文本块列表
    """
    print(f"正在读取文件夹内容: {folder_path}")
    
    if not os.path.exists(folder_path):
        print(f"文件夹不存在: {folder_path}")
        return None
    
    text_blocks = []
    supported_extensions = ['.txt', '.md', '.docx', '.pdf']
    
    try:
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            
            # 跳过目录和隐藏文件
            if os.path.isdir(file_path) or filename.startswith('.'):
                continue
            
            # 检查文件扩展名
            file_ext = os.path.splitext(filename)[1].lower()
            if file_ext not in supported_extensions:
                continue
            
            try:
                if file_ext == '.txt' or file_ext == '.md':
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        if content.strip():
                            # 按段落分割文本块
                            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                            text_blocks.extend(paragraphs)
                            print(f"成功读取文件: {filename}，分割为 {len(paragraphs)} 个文本块")
                # 对于其他格式的文件，可以添加相应的处理逻辑
                
            except Exception as e:
                print(f"读取文件 {filename} 时出错: {e}")
                continue
        
        if text_blocks:
            print(f"成功读取 {len(text_blocks)} 个文本块")
            return text_blocks
        else:
            print("未找到可读取的文本文件")
            return None
            
    except Exception as e:
        print(f"读取文件夹时出错: {e}")
        return None


def generate_exercise_prompt_for_block(content: str, block_index: int, difficulty: str) -> str:
    """
    为单个文本块生成习题的提示词
    """
    difficulty_map = {
        "easy": "简单",
        "medium": "中等", 
        "hard": "困难"
    }
    
    difficulty_text = difficulty_map.get(difficulty, "中等")
    
    prompt = f"""基于以下内容生成一道{difficulty_text}难度的单选题：

{content}

严格按照以下格式生成，不要添加任何其他内容：

题目：
题干：[题干内容]
A. [选项A]
B. [选项B] 
C. [选项C]
D. [选项D]
正确答案：[A/B/C/D]
解析：[详细解析为什么选择正确答案]
所属知识点：[所属知识点]"""

    return prompt


async def generate_exercises_for_blocks(text_blocks: List[str], request: Request, max_tokens: int, temperature: float, difficulty: str) -> List[str]:
    """
    为每个文本块生成一道习题
    """
    print(f"开始为 {len(text_blocks)} 个文本块生成习题...")
    
    exercises = []
    
    for i, block in enumerate(text_blocks):
        print(f"正在为文本块 {i+1}/{len(text_blocks)} 生成习题...")
        print(f"文本块长度: {len(block)} 字符")
        print(f"文本块预览: {block[:200]}...")
        
        try:
            # 为当前文本块生成提示词
            prompt = generate_exercise_prompt_for_block(block, i, difficulty)
            print(f"生成的提示词长度: {len(prompt)} 字符")
            
            # 生成习题
            exercise = await generate_exercises_with_rwkv(
                prompt, 
                request, 
                max_tokens, 
                temperature
            )
            
            # 验证生成的习题
            if not exercise or len(exercise.strip()) == 0:
                print(f"⚠️ 警告：文本块 {i+1} 生成的习题为空")
                exercise = f"题目{i+1}生成失败: 生成的内容为空"
            
            exercises.append(exercise)
            print(f"文本块 {i+1} 习题生成完成，长度: {len(exercise)} 字符")
            
        except Exception as e:
            print(f"为文本块 {i+1} 生成习题时出错: {e}")
            import traceback
            traceback.print_exc()
            # 如果某个文本块生成失败，添加错误信息
            exercises.append(f"题目{i+1}生成失败: {str(e)}")
    
    # 验证生成的习题列表
    if not exercises or len(exercises) == 0:
        print("⚠️ 错误：没有生成任何习题")
        raise Exception("没有生成任何习题")
    
    print(f"成功生成 {len(exercises)} 道习题")
    return exercises


def generate_exercise_prompt(content: str, question_count: int, difficulty: str) -> str:
    """
    生成习题的提示词
    """
    difficulty_map = {
        "easy": "简单",
        "medium": "中等", 
        "hard": "困难"
    }
    
    difficulty_text = difficulty_map.get(difficulty, "中等")
    
    prompt = f"""基于以下内容生成{question_count}道{difficulty_text}难度的单选题：

{content}

严格按照以下格式生成，不要添加任何其他内容：

题目1：
题干：[题干内容]
A. [选项A]
B. [选项B]
C. [选项C]
D. [选项D]
正确答案：[A/B/C/D]
解析：[详细解析为什么选择正确答案]
所属知识点：[所属知识点]

题目2：
题干：[题干内容]
A. [选项A]
B. [选项B]
C. [选项C]
D. [选项D]
正确答案：[A/B/C/D]
解析：[详细解析为什么选择正确答案]
所属知识点：[所属知识点]

[继续生成剩余题目，必须生成{question_count}道题目，每道题都要完整包含题干、选项、正确答案、解析和知识点]"""

    return prompt


def save_exercises_to_file(user_id: str, course_id: str, lesson_num: str, exercises: List[Dict], is_teacher: bool) -> dict:
    """
    将生成的习题保存到文件
    """
    print(f"正在保存习题到文件: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}")
    
    # 获取用户路径
    user_path = get_user_path(user_id, is_teacher)
    
    # 构建保存路径
    lesson_path = os.path.join(user_path, course_id, lesson_num)
    exercises_dir = os.path.join(lesson_path, "exercises")
    
    # 创建习题目录
    os.makedirs(exercises_dir, exist_ok=True)
    
    # 生成文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"exercises_{timestamp}.json"
    file_path = os.path.join(lesson_path, filename)
    
    try:
        # 准备保存的数据
        save_data = {
            "metadata": {
                "user_id": user_id,
                "course_id": course_id,
                "lesson_num": lesson_num,
                "is_teacher": is_teacher,
                "generated_at": datetime.now().isoformat(),
                "total_questions": len(exercises),
                "difficulty": exercises[0]["difficulty"] if exercises else "medium"
            },
            "exercises": exercises
        }
        
        # 保存到文件
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)
        
        print(f"习题已保存到: {file_path}")
        
        return {
            "success": True,
            "file_path": file_path,
            "filename": filename,
            "total_questions": len(exercises)
        }
        
    except Exception as e:
        print(f"保存习题文件时出错: {e}")
        return {
            "success": False,
            "error": str(e)
        }


def save_exercises_to_docx(user_id: str, course_id: str, lesson_num: str, exercise_content: str, is_teacher: bool) -> dict:
    """
    将生成的习题保存为docx文件
    """
    print(f"正在保存习题到docx文件: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}")
    print(f"习题内容长度: {len(exercise_content) if exercise_content else 0} 字符")
    
    # 验证内容是否为空
    if not exercise_content or len(exercise_content.strip()) == 0:
        print("⚠️ 错误：习题内容为空，无法保存")
        return {
            "success": False,
            "error": "习题内容为空",
            "message": "生成的习题内容为空，请检查模型生成是否正常"
        }
    
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建保存路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        
        # 确保目录存在
        os.makedirs(exercises_dir, exist_ok=True)
        
        # 生成文件名：课程号_课时号_时间戳.docx
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{course_id}_{lesson_num}_{timestamp}.docx"
        file_path = os.path.join(exercises_dir, filename)
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        title = doc.add_heading(f'课程{course_id} 课时{lesson_num} 习题', 0)
        title.alignment = 1  # 居中对齐
        
        # 添加生成时间
        time_paragraph = doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        time_paragraph.alignment = 1  # 居中对齐
        
        # 添加分隔线
        doc.add_paragraph('=' * 50)
        
        # 确保内容不为空
        if exercise_content and len(exercise_content.strip()) > 0:
            # 按行分割内容
            lines = exercise_content.split('\n')
            print(f"处理 {len(lines)} 行内容...")
            
            for i, line in enumerate(lines):
                line = line.strip()
                if line:  # 只添加非空行
                    # 简单的内容分类
                    if line.startswith('题目') or (line.startswith('第') and '题' in line):
                        doc.add_heading(line, level=1)
                        print(f"添加标题: {line[:50]}...")
                    elif line.startswith(('A.', 'B.', 'C.', 'D.')):
                        doc.add_paragraph(line)
                        print(f"添加选项: {line[:50]}...")
                    elif line.startswith('正确答案：') or line.startswith('解析：') or line.startswith('知识点：'):
                        p = doc.add_paragraph(line)
                        p.runs[0].bold = True
                        print(f"添加加粗内容: {line[:50]}...")
                    else:
                        doc.add_paragraph(line)
                        print(f"添加普通内容: {line[:50]}...")
            
            print(f"已添加 {len(lines)} 行内容到文档")
        else:
            print("⚠️ 警告：内容为空，添加默认内容")
            doc.add_paragraph("习题内容为空，请检查生成过程")
        
        # 保存文档
        print(f"正在保存文档到: {file_path}")
        try:
            doc.save(file_path)
            print("文档保存成功")
        except Exception as save_error:
            print(f"保存文档时出错: {save_error}")
            raise save_error
        
        # 验证文件是否成功保存且不为空
        if os.path.exists(file_path):
            file_size = os.path.getsize(file_path)
            print(f"文件保存成功，大小: {file_size} 字节")
            
            if file_size == 0:
                print("⚠️ 警告：保存的文件大小为0字节")
                print("尝试使用备用方法保存...")
                
                # 备用方法：直接写入文本文件
                try:
                    txt_filename = filename.replace('.docx', '.txt')
                    txt_file_path = os.path.join(exercises_dir, txt_filename)
                    
                    with open(txt_file_path, 'w', encoding='utf-8') as f:
                        f.write(f"课程{course_id} 课时{lesson_num} 习题\n")
                        f.write(f"生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M:%S')}\n")
                        f.write("=" * 50 + "\n")
                        f.write(exercise_content)
                    
                    txt_file_size = os.path.getsize(txt_file_path)
                    print(f"备用文件保存成功，大小: {txt_file_size} 字节")
                    
                    return {
                        "success": True,
                        "file_path": txt_file_path,
                        "filename": txt_filename,
                        "message": f"习题已保存为 {txt_filename} (备用格式)",
                        "file_size": txt_file_size
                    }
                    
                except Exception as e:
                    print(f"备用保存也失败: {e}")
                    return {
                        "success": False,
                        "error": "文件保存失败",
                        "message": f"无法保存文件: {str(e)}"
                    }
        else:
            print("⚠️ 错误：文件保存失败")
            return {
                "success": False,
                "error": "文件保存失败",
                "message": "无法保存文件到指定路径"
            }
        
        print(f"习题已保存到docx文件: {file_path}")
        
        return {
            "success": True,
            "file_path": file_path,
            "filename": filename,
            "message": f"习题已保存为 {filename}",
            "file_size": file_size
        }
        
    except Exception as e:
        print(f"保存docx文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "message": f"保存docx文件失败: {str(e)}"
        }


async def generate_exercises_with_rwkv(prompt: str, request: Request, max_tokens: int, temperature: float):
    """
    使用RWKV模型生成习题
    """
    print("开始使用RWKV模型生成习题...")
    print(f"提示词长度: {len(prompt)} 字符")
    print(f"提示词预览: {prompt[:200]}...")
    
    try:
        # 获取RWKV模型实例
        model = global_var.get(global_var.Model)
        if model is None:
            raise Exception("RWKV模型未初始化")
        
        # 设置生成参数
        model.temperature = temperature
        model.top_p = 0.9
        
        # 重要：临时增加max_tokens_per_generation以支持更长的生成
        original_max_tokens = model.max_tokens_per_generation
        model.max_tokens_per_generation = max_tokens
        
        # 生成回答内容
        answer_content = ""
        token_count = 0
        
        print("开始生成习题...")
        print(f"最大token数: {max_tokens}")
        print(f"温度参数: {temperature}")
        print(f"模型原始max_tokens_per_generation: {original_max_tokens}")
        print(f"临时设置为: {model.max_tokens_per_generation}")
        
        # 移除所有停止条件，让模型完整生成
        stop_sequences = []
        
        print("开始生成循环...")
        generation_start_time = datetime.now()
        
        for response, delta, _, _ in model.generate(prompt, stop=stop_sequences):
            answer_content += delta
            token_count += 1
            
            # 每100个token打印一次进度
            if token_count % 100 == 0:
                print(f"已生成 {token_count} tokens, 当前内容长度: {len(answer_content)} 字符")
            
            # 检查请求是否断开
            if await request.is_disconnected():
                print("请求已断开")
                break
            
            # 检查是否达到最大token数
            if token_count >= max_tokens:
                print(f"达到最大token数: {max_tokens}")
                break
        
        # 恢复原始设置
        model.max_tokens_per_generation = original_max_tokens
        
        generation_end_time = datetime.now()
        generation_duration = (generation_end_time - generation_start_time).total_seconds()
        
        print(f"习题生成完成，生成长度: {len(answer_content)} 字符")
        print(f"实际生成token数: {token_count}")
        print(f"生成耗时: {generation_duration:.2f}秒")
        print(f"内容预览: {answer_content[:500]}...")
        
        # 检查生成的内容是否为空
        if not answer_content or len(answer_content.strip()) == 0:
            print("⚠️ 错误：RWKV模型生成的内容为空")
            raise Exception("模型生成的内容为空，请检查模型状态和参数设置")
        
        # 检查生成的内容是否太短
        if len(answer_content.strip()) < 50:
            print("⚠️ 警告：生成的内容过短，可能存在问题")
            print(f"内容长度: {len(answer_content)} 字符")
            print(f"内容: {answer_content}")
            
            # 如果内容太短，尝试重新生成
            if token_count < max_tokens * 0.5:  # 如果只用了不到一半的token
                print("尝试继续生成更多内容...")
                # 这里可以添加重试逻辑，但为了简单起见，我们继续使用当前内容
        
        # 清理生成的内容（可选，如果清理后内容为空则使用原始内容）
        try:
            cleaned_content = clean_generated_content(answer_content)
            
            # 检查清理后的内容
            if not cleaned_content or len(cleaned_content.strip()) == 0:
                print("⚠️ 警告：清理后内容为空，使用原始内容")
                cleaned_content = answer_content.strip()
            
            # 如果清理后的内容太短，使用原始内容
            if len(cleaned_content.strip()) < 30:  # 降低阈值
                print("⚠️ 警告：清理后内容过短，使用原始内容")
                cleaned_content = answer_content.strip()
                
        except Exception as e:
            print(f"清理内容时出错: {e}")
            print("使用原始内容")
            cleaned_content = answer_content.strip()
        
        # 最终验证：确保返回的内容不为空
        if not cleaned_content or len(cleaned_content.strip()) == 0:
            print("⚠️ 错误：最终内容为空，使用原始内容")
            cleaned_content = answer_content.strip()
        
        print(f"最终返回内容长度: {len(cleaned_content)} 字符")
        print(f"最终内容预览: {cleaned_content[:300]}...")
        
        return cleaned_content
        
    except Exception as e:
        print(f"RWKV生成习题时出错: {e}")
        import traceback
        traceback.print_exc()
        
        # 确保在出错时也恢复原始设置
        try:
            model = global_var.get(global_var.Model)
            if model:
                model.max_tokens_per_generation = original_max_tokens
        except:
            pass
        raise e


def clean_generated_content(content: str) -> str:
    """
    清理生成的内容，移除不相关的部分
    """
    print("开始清理生成内容...")
    print(f"原始内容长度: {len(content)} 字符")
    print(f"原始内容预览: {content[:500]}...")
    
    if not content or len(content.strip()) == 0:
        print("⚠️ 警告：原始内容为空")
        return ""
    
    # 更保守的清理策略 - 只移除明显无关的前缀
    prefixes_to_remove = [
        "Assistant:",
        "好的，我会根据题目要求继续生成",
        "好的，我可以为您生成",
        "好的，我会为您生成",
    ]
    
    cleaned = content.strip()
    
    # 移除前缀
    for prefix in prefixes_to_remove:
        if cleaned.startswith(prefix):
            cleaned = cleaned[len(prefix):].strip()
            print(f"移除了前缀: {prefix}")
    
    # 如果清理后内容为空，返回原始内容
    if not cleaned or len(cleaned.strip()) == 0:
        print("⚠️ 警告：清理后内容为空，返回原始内容")
        return content.strip()
    
    # 更保守的内容过滤逻辑 - 只过滤明显无关的内容
    lines = cleaned.split('\n')
    filtered_lines = []
    
    for line in lines:
        line = line.strip()
        
        # 只跳过明显无关的内容，避免误删有用内容
        skip_keywords = [
            'Question:', 'Answer:', 'Human:', 'User:', 'Bot:',
            'FIFO', 'SJF', '轮询调度器'
        ]
        
        should_skip = any(keyword in line for keyword in skip_keywords)
        
        if should_skip:
            print(f"跳过包含关键词的行: {line[:50]}...")
            continue
        
        # 保留所有其他内容
        filtered_lines.append(line)
    
    cleaned = '\n'.join(filtered_lines)
    
    # 如果过滤后内容为空，返回清理前的内容
    if not cleaned or len(cleaned.strip()) == 0:
        print("⚠️ 警告：过滤后内容为空，返回清理前的内容")
        return content.strip()
    
    # 最终验证：确保清理后的内容仍然有意义
    if len(cleaned.strip()) < 20:  # 进一步降低阈值
        print("⚠️ 警告：清理后内容过短，可能过度清理，返回原始内容")
        return content.strip()
    
    print(f"清理后内容长度: {len(cleaned)} 字符")
    print(f"清理后内容预览: {cleaned[:500]}...")
    
    return cleaned.strip()


async def generate_with_rwkv_retry(prompt: str, request: Request, max_tokens: int, model, generation_config):
    """
    带重试机制的RWKV生成
    """
    max_retries = 3
    retry_count = 0
    
    while retry_count < max_retries:
        try:
            print(f"RWKV生成尝试 {retry_count + 1}/{max_retries}")
            
            # 设置模型参数
            model.temperature = generation_config.get("temperature", 0.7)
            model.top_p = generation_config.get("top_p", 0.9)
            
            # 重要：临时增加max_tokens_per_generation以支持更长的生成
            original_max_tokens = model.max_tokens_per_generation
            model.max_tokens_per_generation = max_tokens
            
            # 生成回答内容
            answer_content = ""
            token_count = 0
            
            # 停止条件
            stop_sequences = generation_config.get("stop", ["\n\n", "```", "---"])
            
            for response, delta, _, _ in model.generate(prompt, stop=stop_sequences):
                answer_content += delta
                token_count += 1
                
                # 检查请求是否断开
                if await request.is_disconnected():
                    print("请求已断开")
                    break
                
                # 检查是否达到最大token数
                if token_count >= max_tokens:
                    print(f"达到最大token数: {max_tokens}")
                    break
            
            # 恢复原始设置
            model.max_tokens_per_generation = original_max_tokens
            
            # 尝试解析JSON
            try:
                # 查找JSON内容
                start_idx = answer_content.find('{')
                end_idx = answer_content.rfind('}') + 1
                
                if start_idx != -1 and end_idx > start_idx:
                    json_str = answer_content[start_idx:end_idx]
                    json.loads(json_str)  # 验证JSON格式
                    print("JSON格式验证成功")
                    return {"content": answer_content}
                else:
                    print("未找到有效的JSON内容")
                    raise Exception("JSON格式无效")
                    
            except json.JSONDecodeError as json_error:
                print(f"JSON解析失败: {json_error}")
                if retry_count < max_retries - 1:
                    retry_count += 1
                    continue
                else:
                    raise Exception(f"JSON格式错误: {json_error}")
            
        except Exception as e:
            print(f"RWKV生成出错: {e}")
            if retry_count < max_retries - 1:
                retry_count += 1
                await asyncio.sleep(1)  # 等待1秒后重试
                continue
            else:
                raise e
    
    raise Exception("RWKV生成失败，已达到最大重试次数")


def parse_exercises_from_response(response_text: str) -> List[Dict]:
    """
    从RWKV响应中解析习题数据
    """
    print("正在解析习题响应...")
    print(f"原始响应内容: {response_text}")
    
    try:
        # 查找JSON内容
        start_idx = response_text.find('{')
        end_idx = response_text.rfind('}') + 1
        
        if start_idx == -1 or end_idx <= start_idx:
            raise Exception("未找到有效的JSON内容")
        
        json_str = response_text[start_idx:end_idx]
        print(f"提取的JSON字符串: {json_str}")
        
        # 尝试清理JSON字符串
        json_str = clean_json_string(json_str)
        print(f"清理后的JSON字符串: {json_str}")
        
        data = json.loads(json_str)
        
        if "questions" not in data:
            raise Exception("响应中缺少questions字段")
        
        questions = data["questions"]
        if not isinstance(questions, list):
            raise Exception("questions字段不是列表格式")
        
        # 验证每个题目的格式
        parsed_questions = []
        for i, question in enumerate(questions):
            if not isinstance(question, dict):
                print(f"跳过无效题目 {i}: 不是字典格式")
                continue
            
            # 检查必需字段
            required_fields = ["question_text", "options", "correct_answer", "explanation", "knowledge_point"]
            missing_fields = [field for field in required_fields if field not in question]
            
            if missing_fields:
                print(f"跳过题目 {i}: 缺少字段 {missing_fields}")
                continue
            
            # 生成题目ID
            question_id = question.get("question_id", f"q{i+1}")
            
            # 验证选项格式
            options = question["options"]
            if not isinstance(options, list) or len(options) != 4:
                print(f"跳过题目 {i}: 选项格式不正确")
                continue
            
            # 验证正确答案
            correct_answer = question["correct_answer"]
            if correct_answer not in ["A", "B", "C", "D"]:
                print(f"跳过题目 {i}: 正确答案格式不正确")
                continue
            
            # 构建标准化的题目对象
            parsed_question = {
                "question_id": question_id,
                "question_text": question["question_text"],
                "options": options,
                "correct_answer": correct_answer,
                "explanation": question["explanation"],
                "knowledge_point": question["knowledge_point"],
                "difficulty": question.get("difficulty", "medium")
            }
            
            parsed_questions.append(parsed_question)
            print(f"成功解析题目 {i+1}: {question_id}")
        
        print(f"成功解析 {len(parsed_questions)} 道题目")
        return parsed_questions
        
    except Exception as e:
        print(f"解析习题响应时出错: {e}")
        raise e


def clean_json_string(json_str: str) -> str:
    """
    清理JSON字符串，修复常见的格式问题
    """
    # 移除可能的BOM标记和多余空白
    json_str = json_str.strip()
    
    # 修复可能的换行符问题
    json_str = json_str.replace('\n', '\\n').replace('\r', '\\r')
    
    # 修复可能的制表符问题
    json_str = json_str.replace('\t', '\\t')
    
    # 修复多余的逗号（在数组或对象结束前）
    json_str = json_str.replace(',}', '}').replace(',]', ']')
    
    return json_str


@router.post("/v1/exercise/generate", tags=["Exercise"], response_model=ExerciseResponse)
async def generate_exercises(body: ExerciseBody, request: Request):
    """
    根据课时知识库内容生成课后习题，支持两种生成模式
    """
    start_time = datetime.now()
    
    print(f"开始生成习题: userID={body.user_id}, courseId={body.course_id}, lessonNum={body.lesson_num}")
    print(f"参数: questionCount={body.question_count}, difficulty={body.difficulty}, generationMode={body.generation_mode}")
    
    # 检查并发锁
    if exercise_lock.locked():
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="系统繁忙，请稍后重试"
        )
    
    try:
        with exercise_lock:
            if body.generation_mode == "whole":
                # 整体内容生成模式
                print("使用整体内容生成模式")
                
                # 1. 获取合并的课时内容
                print("步骤1: 获取课时内容")
                lesson_content = get_lesson_content_whole(
                    body.user_id, 
                    body.course_id, 
                    body.lesson_num, 
                    body.is_teacher
                )
                
                if not lesson_content:
                    raise HTTPException(
                        status_code=status.HTTP_404_NOT_FOUND,
                        detail="无法获取课时内容，请确保课时已上传并处理完成"
                    )
                
                print(f"课时内容长度: {len(lesson_content)} 字符")
                
                # 2. 生成提示词
                print("步骤2: 生成提示词")
                prompt = generate_exercise_prompt(
                    lesson_content, 
                    body.question_count, 
                    body.difficulty
                )
                
                # 3. 使用RWKV生成习题
                print("步骤3: 使用RWKV生成习题")
                response_text = await generate_exercises_with_rwkv(
                    prompt, 
                    request, 
                    body.max_tokens, 
                    body.temperature
                )
                
                # 验证生成的内容
                if not response_text or len(response_text.strip()) == 0:
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="习题生成失败：生成的内容为空"
                    )
                
                print(f"生成的内容长度: {len(response_text)} 字符")
                print(f"生成的内容预览: {response_text[:300]}...")
                
                # 4. 保存习题到docx文件
                print("步骤4: 保存习题到docx文件")
                save_result = save_exercises_to_docx(
                    body.user_id,
                    body.course_id,
                    body.lesson_num,
                    response_text,
                    body.is_teacher
                )
                
                if save_result["success"]:
                    print(f"✅ {save_result['message']}")
                    print(f"文件大小: {save_result.get('file_size', 'unknown')} 字节")
                else:
                    print(f"⚠️ 保存docx文件失败: {save_result['error']}")
                    # 即使保存失败，也返回生成的内容
                    print("继续返回生成的内容...")
                
                # 5. 返回结果
                end_time = datetime.now()
                generation_time = (end_time - start_time).total_seconds()
                
                response = ExerciseResponse(
                    success=True,
                    message=f"习题生成完成，共生成 {body.question_count} 道题目",
                    data=response_text,
                    total_count=body.question_count,
                    generation_time=generation_time
                )
                print(f"习题生成完成，耗时: {generation_time:.2f}秒")
                return response
                
            else:
                # 文本块生成模式（默认）
                print("使用文本块生成模式")
                
                # 1. 获取课时内容（文本块列表）
                print("步骤1: 获取课时内容")
                text_blocks = get_lesson_content(
                    body.user_id, 
                    body.course_id, 
                    body.lesson_num, 
                    body.is_teacher
                )
                
                if not text_blocks:
                    raise HTTPException(
                        status_code=status.HTTP_404_NOT_FOUND,
                        detail="无法获取课时内容，请确保课时已上传并处理完成"
                    )
                
                print(f"获取到 {len(text_blocks)} 个文本块")
                
                # 2. 限制文本块数量，避免生成过多题目
                max_blocks = min(body.question_count, len(text_blocks))
                selected_blocks = text_blocks[:max_blocks]
                
                print(f"将使用前 {max_blocks} 个文本块生成习题")
                
                # 3. 为每个文本块生成一道习题
                print("步骤2: 为每个文本块生成习题")
                exercises = await generate_exercises_for_blocks(
                    selected_blocks,
                    request,
                    body.max_tokens,
                    body.temperature,
                    body.difficulty
                )
                
                # 验证生成的习题
                if not exercises or len(exercises) == 0:
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="习题生成失败：没有生成任何习题"
                    )
                
                print(f"成功生成 {len(exercises)} 道习题")
                
                # 4. 合并所有习题
                combined_exercises = "\n\n" + "="*50 + "\n\n".join(exercises)
                
                # 验证合并后的内容
                if not combined_exercises or len(combined_exercises.strip()) == 0:
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail="习题生成失败：合并后的内容为空"
                    )
                
                print(f"合并后内容长度: {len(combined_exercises)} 字符")
                print(f"合并后内容预览: {combined_exercises[:300]}...")
                
                # 5. 保存习题到docx文件
                print("步骤3: 保存习题到docx文件")
                save_result = save_exercises_to_docx(
                    body.user_id,
                    body.course_id,
                    body.lesson_num,
                    combined_exercises,
                    body.is_teacher
                )
                
                if save_result["success"]:
                    print(f"✅ {save_result['message']}")
                    print(f"文件大小: {save_result.get('file_size', 'unknown')} 字节")
                else:
                    print(f"⚠️ 保存docx文件失败: {save_result['error']}")
                    # 即使保存失败，也返回生成的内容
                    print("继续返回生成的内容...")
                
                # 6. 返回结果
                end_time = datetime.now()
                generation_time = (end_time - start_time).total_seconds()
                
                response = ExerciseResponse(
                    success=True,
                    message=f"习题生成完成，共生成 {len(exercises)} 道题目",
                    data=combined_exercises,
                    total_count=len(exercises),
                    generation_time=generation_time
                )
                print(f"习题生成完成，耗时: {generation_time:.2f}秒，生成 {len(exercises)} 道题目")
                return response
            
    except HTTPException:
        raise
    except Exception as e:
        print(f"生成习题时出错: {e}")
        import traceback
        traceback.print_exc()
        
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"生成习题失败: {str(e)}"
        )


@router.get("/v1/exercise/list/{user_id}/{course_id}/{lesson_num}", tags=["Exercise"])
async def get_exercise_list(user_id: str, course_id: str, lesson_num: str, is_teacher: bool = False):
    """
    获取指定课时的习题列表
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建习题目录路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        
        if not os.path.exists(exercises_dir):
            return {
                "success": True,
                "data": [],
                "total": 0,
                "message": "该课时暂无习题"
            }
        
        # 读取所有习题文件（.docx格式）
        exercise_files = []
        for filename in os.listdir(exercises_dir):
            if filename.endswith('.docx'):
                file_path = os.path.join(exercises_dir, filename)
                try:
                    # 获取文件信息
                    file_stat = os.stat(file_path)
                    exercise_files.append({
                        "filename": filename,
                        "file_path": file_path,
                        "file_size": file_stat.st_size,
                        "created_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                        "modified_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                        "file_type": "docx"
                    })
                except Exception as e:
                    print(f"读取习题文件 {filename} 时出错: {e}")
                    continue
        
        # 按修改时间排序（最新的在前）
        exercise_files.sort(key=lambda x: x["modified_time"], reverse=True)
        
        return {
            "success": True,
            "data": exercise_files,
            "total": len(exercise_files),
            "message": f"找到 {len(exercise_files)} 个习题文件"
        }
        
    except Exception as e:
        print(f"获取习题列表时出错: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取习题列表失败: {str(e)}"
        )


@router.get("/v1/exercise/{user_id}/{course_id}/{lesson_num}/{filename}", tags=["Exercise"])
async def get_exercise_detail(user_id: str, course_id: str, lesson_num: str, filename: str, is_teacher: bool = False):
    """
    获取指定习题文件的详细内容
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建文件路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        file_path = os.path.join(exercises_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="习题文件不存在"
            )
        
        # 获取文件信息
        file_stat = os.stat(file_path)
        
        return {
            "success": True,
            "data": {
                "filename": filename,
                "file_path": file_path,
                "file_size": file_stat.st_size,
                "created_time": datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
                "modified_time": datetime.fromtimestamp(file_stat.st_mtime).isoformat(),
                "file_type": "docx",
                "download_url": f"/v1/exercise/download/{user_id}/{course_id}/{lesson_num}/{filename}"
            },
            "message": "获取习题详情成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"获取习题详情时出错: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取习题详情失败: {str(e)}"
        )


@router.delete("/v1/exercise/{user_id}/{course_id}/{lesson_num}/{filename}", tags=["Exercise"])
async def delete_exercise_file(user_id: str, course_id: str, lesson_num: str, filename: str, is_teacher: bool = False):
    """
    删除指定的习题文件
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建文件路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        file_path = os.path.join(exercises_dir, filename)
        
        if not os.path.exists(file_path):
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="习题文件不存在"
            )
        
        # 删除文件
        os.remove(file_path)
        
        return {
            "success": True,
            "message": "习题文件删除成功"
        }
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"删除习题文件时出错: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"删除习题文件失败: {str(e)}"
        )


@router.get("/v1/exercise/download/{user_id}/{course_id}/{lesson_num}/{filename}", tags=["Exercise"])
async def download_exercise_file(user_id: str, course_id: str, lesson_num: str, filename: str, is_teacher: bool = False):
    """
    下载指定的习题文件
    """
    try:
        print(f"开始下载文件: userID={user_id}, courseId={course_id}, lessonNum={lesson_num}, filename={filename}")
        
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建文件路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        file_path = os.path.join(exercises_dir, filename)
        
        print(f"文件完整路径: {file_path}")
        
        if not os.path.exists(file_path):
            print(f"文件不存在: {file_path}")
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="习题文件不存在"
            )
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        print(f"文件大小: {file_size} 字节")
        
        if file_size == 0:
            print("⚠️ 警告：文件大小为0字节")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="文件为空，无法下载"
            )
        
        # 读取文件内容
        try:
            with open(file_path, 'rb') as f:
                file_content = f.read()
            print(f"成功读取文件内容，大小: {len(file_content)} 字节")
        except Exception as read_error:
            print(f"读取文件失败: {read_error}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"读取文件失败: {str(read_error)}"
            )
        
        # 验证读取的内容大小
        if len(file_content) != file_size:
            print(f"⚠️ 警告：读取的内容大小({len(file_content)})与文件大小({file_size})不匹配")
        
        # 处理文件名编码
        import urllib.parse
        encoded_filename = urllib.parse.quote(filename)
        
        # 设置响应头
        headers = {
            "Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "Content-Length": str(len(file_content))
        }
        
        print(f"响应头: {headers}")
        print(f"文件内容前100字节: {file_content[:100]}")
        
        return Response(
            content=file_content,
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        print(f"下载习题文件时出错: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"下载习题文件失败: {str(e)}"
        )


@router.get("/v1/exercise/status", tags=["Exercise"])
async def get_exercise_status():
    """
    获取习题生成服务状态
    """
    try:
        # 检查RWKV模型状态
        model = global_var.get(global_var.Model)
        model_status = "ready" if model is not None else "not_ready"
        
        # 获取模型详细信息
        model_info = {}
        if model is not None:
            try:
                model_info = {
                    "temperature": getattr(model, 'temperature', 'unknown'),
                    "top_p": getattr(model, 'top_p', 'unknown'),
                    "max_tokens_per_generation": getattr(model, 'max_tokens_per_generation', 'unknown'),
                    "model_loaded": True
                }
            except Exception as e:
                model_info = {
                    "error": str(e),
                    "model_loaded": False
                }
        
        return {
            "success": True,
            "data": {
                "service": "exercise_generation",
                "status": "running",
                "model_status": model_status,
                "model_info": model_info,
                "timestamp": datetime.now().isoformat()
            },
            "message": "习题生成服务运行正常"
        }
        
    except Exception as e:
        print(f"获取服务状态时出错: {e}")
        return {
            "success": False,
            "data": {
                "service": "exercise_generation",
                "status": "error",
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            },
            "message": "习题生成服务异常"
        }


@router.get("/v1/exercise/debug", tags=["Exercise"])
async def debug_exercise_generation():
    """
    调试习题生成功能
    """
    try:
        # 检查RWKV模型状态
        model = global_var.get(global_var.Model)
        
        debug_info = {
            "model_available": model is not None,
            "model_type": type(model).__name__ if model else None,
            "global_var_available": hasattr(global_var, 'Model'),
            "settings_available": True,  # 假设设置总是可用的
        }
        
        if model is not None:
            try:
                # 测试模型基本功能
                debug_info.update({
                    "temperature": getattr(model, 'temperature', 'unknown'),
                    "top_p": getattr(model, 'top_p', 'unknown'),
                    "max_tokens_per_generation": getattr(model, 'max_tokens_per_generation', 'unknown'),
                    "model_attributes": [attr for attr in dir(model) if not attr.startswith('_')]
                })
            except Exception as e:
                debug_info["model_error"] = str(e)
        
        # 检查文件系统权限
        try:
            test_dir = "/tmp/exercise_test"
            os.makedirs(test_dir, exist_ok=True)
            test_file = os.path.join(test_dir, "test.txt")
            with open(test_file, 'w') as f:
                f.write("test")
            os.remove(test_file)
            os.rmdir(test_dir)
            debug_info["file_system_ok"] = True
        except Exception as e:
            debug_info["file_system_error"] = str(e)
            debug_info["file_system_ok"] = False
        
        return {
            "success": True,
            "data": debug_info,
            "message": "调试信息获取成功"
        }
        
    except Exception as e:
        print(f"获取调试信息时出错: {e}")
        return {
            "success": False,
            "data": {
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            },
            "message": "获取调试信息失败"
        }


@router.post("/v1/exercise/test", tags=["Exercise"])
async def test_exercise_generation():
    """
    测试习题生成功能
    """
    try:
        # 检查RWKV模型状态
        model = global_var.get(global_var.Model)
        if model is None:
            return {
                "success": False,
                "error": "RWKV模型未初始化",
                "message": "请确保模型已正确加载"
            }
        
        # 创建测试内容
        test_content = """
        这是一个测试内容，用于验证习题生成功能。
        这里包含一些基本的概念和知识点，可以用来生成习题。
        测试内容包括：
        1. 基本概念
        2. 重要知识点
        3. 应用场景
        """
        
        # 生成测试提示词
        test_prompt = generate_exercise_prompt(test_content, 1, "medium")
        
        # 模拟请求对象
        class MockRequest:
            async def is_disconnected(self):
                return False
        
        mock_request = MockRequest()
        
        # 测试生成
        try:
            result = await generate_exercises_with_rwkv(
                test_prompt,
                mock_request,
                1000,  # max_tokens
                0.7    # temperature
            )
            
            if result and len(result.strip()) > 0:
                return {
                    "success": True,
                    "data": {
                        "generated_content": result[:500] + "..." if len(result) > 500 else result,
                        "content_length": len(result),
                        "model_working": True
                    },
                    "message": "习题生成测试成功"
                }
            else:
                return {
                    "success": False,
                    "error": "生成的内容为空",
                    "message": "模型生成的内容为空，请检查模型状态"
                }
                
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "message": f"习题生成测试失败: {str(e)}"
            }
        
    except Exception as e:
        print(f"测试习题生成时出错: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "data": {
                "error": str(e),
                "timestamp": datetime.now().isoformat()
            },
            "message": "测试习题生成失败"
        }


@router.post("/v1/exercise/debug-save", tags=["Exercise"])
async def debug_save_exercise():
    """
    调试文件保存功能
    """
    try:
        # 测试内容
        test_content = """
题目1：什么是操作系统？
A. 计算机硬件
B. 管理计算机硬件和软件资源的系统软件
C. 应用程序
D. 网络设备

正确答案：B

解析：操作系统是管理计算机硬件和软件资源的系统软件，它为用户和其他软件提供公共服务。

题目2：进程和线程的区别是什么？
A. 没有区别
B. 进程是资源分配的基本单位，线程是CPU调度的基本单位
C. 进程比线程更轻量级
D. 线程比进程更重

正确答案：B

解析：进程是资源分配的基本单位，拥有独立的内存空间；线程是CPU调度的基本单位，共享进程的资源。
"""
        
        # 测试保存
        save_result = save_exercises_to_docx(
            "test_user",
            "test_course", 
            "test_lesson",
            test_content,
            False
        )
        
        return {
            "success": True,
            "save_result": save_result,
            "test_content_length": len(test_content),
            "message": "调试保存完成"
        }
        
    except Exception as e:
        print(f"调试保存失败: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "message": f"调试保存失败: {str(e)}"
        }


@router.get("/v1/exercise/test-docx/{user_id}/{course_id}/{lesson_num}/{filename}", tags=["Exercise"])
async def test_docx_file(user_id: str, course_id: str, lesson_num: str, filename: str, is_teacher: bool = False):
    """
    测试docx文件是否可以正常打开
    """
    try:
        # 获取用户路径
        user_path = get_user_path(user_id, is_teacher)
        
        # 构建文件路径
        lesson_path = os.path.join(user_path, course_id, lesson_num)
        exercises_dir = os.path.join(lesson_path, "exercises")
        file_path = os.path.join(exercises_dir, filename)
        
        print(f"测试文件路径: {file_path}")
        
        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": "文件不存在",
                "file_path": file_path
            }
        
        # 检查文件大小
        file_size = os.path.getsize(file_path)
        print(f"文件大小: {file_size} 字节")
        
        # 尝试读取文件内容
        try:
            with open(file_path, 'rb') as f:
                content = f.read(1000)  # 读取前1000字节
            content_readable = True
            content_length = len(content)
            print(f"文件可读，读取了 {content_length} 字节")
        except Exception as read_error:
            content_readable = False
            content_length = 0
            print(f"文件读取失败: {read_error}")
        
        # 尝试用python-docx打开文件
        try:
            test_doc = Document(file_path)
            paragraphs = len(test_doc.paragraphs)
            print(f"docx文件可以正常打开，包含 {paragraphs} 个段落")
            docx_readable = True
        except Exception as docx_error:
            print(f"docx文件无法打开: {docx_error}")
            docx_readable = False
            paragraphs = 0
        
        return {
            "success": True,
            "data": {
                "filename": filename,
                "file_path": file_path,
                "file_size": file_size,
                "content_readable": content_readable,
                "content_length": content_length,
                "docx_readable": docx_readable,
                "paragraphs": paragraphs
            },
            "message": "文件测试完成"
        }
        
    except Exception as e:
        print(f"测试docx文件时出错: {e}")
        import traceback
        traceback.print_exc()
        return {
            "success": False,
            "error": str(e),
            "message": f"测试docx文件失败: {str(e)}"
        }
